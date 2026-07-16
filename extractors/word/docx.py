from __future__ import annotations

from pathlib import Path

from io import BytesIO

from docx import Document as DocxDocument

from document_engine.models import Document, ImageElement, TableElement
from document_engine.core.registry import Extractor, register_extractor


class DocxExtractor(Extractor):
    format = "docx"

    def extract_from_stream(self, stream: bytes, filename: str = "upload.docx") -> Document:
        docx = DocxDocument(BytesIO(stream))
        document = Document(
            title=Path(filename).stem,
            file_path=filename,
            file_format="docx",
            file_size_bytes=len(stream),
        )
        return self._process_docx(docx, document)

    def extract(self, file_path: str) -> Document:
        docx = DocxDocument(file_path)
        p = Path(file_path)

        document = Document(
            title=p.stem,
            file_path=file_path,
            file_format="docx",
            file_size_bytes=p.stat().st_size,
        )

        return self._process_docx(docx, document)

    def _process_docx(self, docx: DocxDocument, document: Document) -> Document:
        paragraphs: list[str] = []
        for para in docx.paragraphs:
            paragraphs.append(para.text)
        document.text = "\n".join(paragraphs)

        for table in docx.tables:
            headers = [cell.text for cell in table.rows[0].cells] if table.rows else []
            rows = [[cell.text for cell in row.cells] for row in table.rows[1:]]
            document.tables.append(TableElement(headers=headers, rows=rows))

        document.raw_metadata["paragraph_count"] = len(docx.paragraphs)
        document.raw_metadata["table_count"] = len(docx.tables)
        document.raw_metadata["section_count"] = len(docx.sections)

        return document

    def extract_text(self, file_path: str) -> str:
        docx = DocxDocument(file_path)
        return "\n".join(p.text for p in docx.paragraphs)


register_extractor(DocxExtractor())
